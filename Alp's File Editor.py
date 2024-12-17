import os
import subprocess
import sys

# Function to check and install a package
def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# Function to run terminal commands
def run_command(command):
    try:
        subprocess.check_call(command, shell=True)
    except subprocess.CalledProcessError as e:
        print(f"Error while running command: {command}\nError: {e}")

# Install required Python packages
required_packages = ['Pillow', 'PyMuPDF', 'python-docx', 'moviepy','PyPDF2']
for package in required_packages:
    try:
        __import__(package)
    except ImportError:
        print(f"{package} is not installed. Installing...")
        try:
            install(package)
            print(f"{package} installed successfully.")
        except Exception as e:
            print(f"Failed to install {package}. Error: {str(e)}")



print("All required modules and dependencies have been installed and configured.")

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image
from moviepy.editor import VideoFileClip
from PyPDF2 import PdfMerger
from PyPDF2.errors import PdfReadError  
import fitz  
from docx import Document


video_formats = ["MP4", "MOV", "AVI", "FLV", "MKV", "WMV", "AVCHD", "WEBM", "MPEG", "3GP", "RMVB", "VOB", "TS", "MPG", "M4V"]
image_formats = ["JPEG", "PNG", "BMP", "GIF", "TIFF", "WEBP", "ICO", "TGA", "PCX", "JP2"]

def convert_image(input_file, output_format):
    try:
        img = Image.open(input_file)
        
        directory, filename = os.path.split(input_file)
        filename_without_extension = os.path.splitext(filename)[0]
        output_file = os.path.join(directory, filename_without_extension + '.' + output_format.lower())
        
        img.save(output_file)
        messagebox.showinfo("Success", f"Conversion successful. Image saved at: {output_file}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def convert_video(input_file, output_format):
    try:
        clip = VideoFileClip(input_file)
        
        directory, filename = os.path.split(input_file)
        filename_without_extension = os.path.splitext(filename)[0]
        output_file = os.path.join(directory, filename_without_extension + '.' + output_format.lower())
        
        clip.write_videofile(output_file, codec='libx264')  # Fix: Ensure libx264 codec for common formats
        messagebox.showinfo("Success", f"Conversion successful. Video saved at: {output_file}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def convert_pdf_to_docx(pdf_file, docx_file):
    try:
        doc = fitz.open(pdf_file)  # PyMuPDF opens the PDF
        text = ""
        for page in doc:
            text += page.get_text()
        
        word_doc = Document()  # Fix: Use Document for saving DOCX
        word_doc.add_paragraph(text)
        word_doc.save(docx_file)  # Save DOCX file
        
        doc.close()
        messagebox.showinfo("Success", f"Conversion successful. DOCX file saved at: {docx_file}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")


def combine_docx_files(num_files):
    def browse_file(index):
        file_path = filedialog.askopenfilename(title=f"Select DOCX File {index}", filetypes=[("DOCX files", "*.docx")])
        if file_path:
            entry_paths[index - 1].delete(0, tk.END)
            entry_paths[index - 1].insert(0, file_path)

    def combine():
        combined_doc = Document()
        for entry in entry_paths:
            file_path = entry.get()
            if file_path:
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    combined_doc.add_paragraph(paragraph.text)
                combined_doc.add_page_break()

        combined_file_name = entry_combined_file_name.get()
        combined_doc.save(combined_file_name)
        messagebox.showinfo("Success", f"DOCX files combined successfully into: {combined_file_name}")

    window = tk.Toplevel()
    window.title("Combine DOCX Files")

    label_instructions = tk.Label(window, text=f"Enter the number of docx files you want to combine: {num_files}")
    label_instructions.grid(row=0, column=0, columnspan=2, padx=5, pady=5)

    entry_paths = []
    for i in range(1, num_files + 1):
        label_file = tk.Label(window, text=f"DOCX File {i}")
        label_file.grid(row=i, column=0, padx=5, pady=5)

        entry_path = tk.Entry(window, width=50)
        entry_path.grid(row=i, column=1, padx=5, pady=5)

        button_browse = tk.Button(window, text="Browse", command=lambda idx=i: browse_file(idx))
        button_browse.grid(row=i, column=2, padx=5, pady=5)

        entry_paths.append(entry_path)

    label_combined_file_name = tk.Label(window, text="Enter the name of the combined docx file:")
    label_combined_file_name.grid(row=num_files + 1, column=0, padx=5, pady=5)

    entry_combined_file_name = tk.Entry(window, width=50)
    entry_combined_file_name.grid(row=num_files + 1, column=1, columnspan=2, padx=5, pady=5)

    button_combine = tk.Button(window, text="Combine", command=combine)
    button_combine.grid(row=num_files + 2, column=1, columnspan=2, padx=5, pady=5)

def combine_pdf_files(num_files):
    def browse_file(index):
        file_path = filedialog.askopenfilename(title=f"Select PDF File {index}", filetypes=[("PDF files", "*.pdf")])
        if file_path:
            entry_paths[index - 1].delete(0, tk.END)
            entry_paths[index - 1].insert(0, file_path)

    def combine():
        merger = PdfMerger()
        for entry in entry_paths:
            file_path = entry.get()
            if file_path:
                merger.append(file_path)

        combined_file_name = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if combined_file_name:
            with open(combined_file_name, "wb") as f:
                merger.write(f)
            messagebox.showinfo("Success", f"PDFs combined successfully. Combined PDF saved at: {combined_file_name}")

    window = tk.Toplevel()
    window.title("Combine PDF Files")

    label_instructions = tk.Label(window, text=f"Enter the number of PDF files you want to combine: {num_files}")
    label_instructions.grid(row=0, column=0, columnspan=2, padx=5, pady=5)

    entry_paths = []
    for i in range(1, num_files + 1):
        label_file = tk.Label(window, text=f"PDF File {i}")
        label_file.grid(row=i, column=0, padx=5, pady=5)

        entry_path = tk.Entry(window, width=50)
        entry_path.grid(row=i, column=1, padx=5, pady=5)

        button_browse = tk.Button(window, text="Browse", command=lambda idx=i: browse_file(idx))
        button_browse.grid(row=i, column=2, padx=5, pady=5)

        entry_paths.append(entry_path)

    button_combine = tk.Button(window, text="Combine", command=combine)
    button_combine.grid(row=num_files + 1, column=1, columnspan=2, padx=5, pady=5)

# GUI setup
window = tk.Tk()
window.title("Media Format Converter")

notebook = ttk.Notebook(window)
notebook.pack(fill='both', expand=True)

# Photo tab
photo_tab = ttk.Frame(notebook)
notebook.add(photo_tab, text='Photo Converter')

label_output_format_photo = tk.Label(photo_tab, text="Select output format:")
label_output_format_photo.grid(row=0, column=0, padx=5, pady=5)

output_format_var_photo = tk.StringVar()
output_format_menu_photo = tk.OptionMenu(photo_tab, output_format_var_photo, *image_formats)  # Fix: Populate with image formats
output_format_menu_photo.grid(row=0, column=1, padx=5, pady=5)

label_path_photo = tk.Label(photo_tab, text="Select image file:")
label_path_photo.grid(row=1, column=0, padx=5, pady=5)

entry_path_photo = tk.Entry(photo_tab, width=50)
entry_path_photo.grid(row=1, column=1, padx=5, pady=5)

button_browse_photo = tk.Button(photo_tab, text="Browse", command=lambda: entry_path_photo.insert(0, filedialog.askopenfilename(filetypes=[("Image files", "*.*")])))
button_browse_photo.grid(row=1, column=2, padx=5, pady=5)

button_convert_photo = tk.Button(photo_tab, text="Convert", command=lambda: convert_image(entry_path_photo.get(), output_format_var_photo.get()))
button_convert_photo.grid(row=2, column=1, padx=5, pady=5)

# Video tab
video_tab = ttk.Frame(notebook)
notebook.add(video_tab, text='Video Converter')

label_output_format_video = tk.Label(video_tab, text="Select output format:")
label_output_format_video.grid(row=0, column=0, padx=5, pady=5)

output_format_var_video = tk.StringVar()
output_format_menu_video = tk.OptionMenu(video_tab, output_format_var_video, *video_formats)  # Fix: Populate with video formats
output_format_menu_video.grid(row=0, column=1, padx=5, pady=5)

label_path_video = tk.Label(video_tab, text="Select video file:")
label_path_video.grid(row=1, column=0, padx=5, pady=5)

entry_path_video = tk.Entry(video_tab, width=50)
entry_path_video.grid(row=1, column=1, padx=5, pady=5)

button_browse_video = tk.Button(video_tab, text="Browse", command=lambda: entry_path_video.insert(0, filedialog.askopenfilename(filetypes=[("Video files", "*.*")])))
button_browse_video.grid(row=1, column=2, padx=5, pady=5)

button_convert_video = tk.Button(video_tab, text="Convert", command=lambda: convert_video(entry_path_video.get(), output_format_var_video.get()))
button_convert_video.grid(row=2, column=1, padx=5, pady=5)

# DOCX-PDF tab
pdf_tab = ttk.Frame(notebook)
notebook.add(pdf_tab, text="PDF-DOCX Converter")

label_path_pdf = tk.Label(pdf_tab, text="Select PDF file:")
label_path_pdf.grid(row=0, column=0, padx=5, pady=5)

entry_path_pdf = tk.Entry(pdf_tab, width=50)
entry_path_pdf.grid(row=0, column=1, padx=5, pady=5)

button_browse_pdf = tk.Button(pdf_tab, text="Browse", command=lambda: entry_path_pdf.insert(0, filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])))
button_browse_pdf.grid(row=0, column=2, padx=5, pady=5)

label_output_file_pdf = tk.Label(pdf_tab, text="Enter output DOCX file name:")
label_output_file_pdf.grid(row=1, column=0, padx=5, pady=5)

entry_output_file_pdf = tk.Entry(pdf_tab, width=50)
entry_output_file_pdf.grid(row=1, column=1, padx=5, pady=5)

button_convert_pdf = tk.Button(pdf_tab, text="Convert", command=lambda: convert_pdf_to_docx(entry_path_pdf.get(), entry_output_file_pdf.get()))
button_convert_pdf.grid(row=2, column=1, padx=5, pady=5)

# Combine DOCX tab
combine_docx_tab = ttk.Frame(notebook)
notebook.add(combine_docx_tab, text="Combine DOCX Files")

label_num_files_docx = tk.Label(combine_docx_tab, text="Enter the number of DOCX files to combine:")
label_num_files_docx.grid(row=0, column=0, padx=5, pady=5)

entry_num_files_docx = tk.Entry(combine_docx_tab, width=10)
entry_num_files_docx.grid(row=0, column=1, padx=5, pady=5)

button_combine_docx = tk.Button(combine_docx_tab, text="Combine DOCX Files", command=lambda: combine_docx_files(int(entry_num_files_docx.get())))
button_combine_docx.grid(row=1, column=1, padx=5, pady=5)

# Combine PDF tab
combine_pdf_tab = ttk.Frame(notebook)
notebook.add(combine_pdf_tab, text="Combine PDF Files")

label_num_files_pdf = tk.Label(combine_pdf_tab, text="Enter the number of PDF files to combine:")
label_num_files_pdf.grid(row=0, column=0, padx=5, pady=5)

entry_num_files_pdf = tk.Entry(combine_pdf_tab, width=10)
entry_num_files_pdf.grid(row=0, column=1, padx=5, pady=5)

button_combine_pdf = tk.Button(combine_pdf_tab, text="Combine PDF Files", command=lambda: combine_pdf_files(int(entry_num_files_pdf.get())))
button_combine_pdf.grid(row=1, column=1, padx=5, pady=5)

window.mainloop()

